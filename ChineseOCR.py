import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import cv2
import numpy as np
from docx import Document
import os
import tkinter as tk
from tkinter import filedialog, messagebox
import easyocr

class ChineseOCR:
    def __init__(self):
        self.setup_tesseract_path()
        # 初始化EasyOCR阅读器（支持中文）
        self.easy_reader = easyocr.Reader(['ch_sim', 'en'])  # 使用中文简体和英文

    def setup_tesseract_path(self):
        """设置Tesseract路径"""
        possible_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            r"C:\Users\azzbl\AppData\Roaming\Python\Python311\Scripts\pytesseract.exe"
        ]
        for path in possible_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                print(f"找到Tesseract: {path}")
                return
        print("警告: 未找到Tesseract，将尝试使用EasyOCR")

    def preprocess_image(self, image_path):
        """图像预处理，优化用于中文识别"""
        # 使用OpenCV读取图像
        img = cv2.imread(image_path)
        if img is None:
            # 如果OpenCV无法读取，尝试用PIL
            img = Image.open(image_path)
            img = np.array(img)
            # 如果图像是RGBA，转换为RGB
            if len(img.shape) == 3 and img.shape[2] == 4:
                img = cv2.cvtColor(img, cv2.COLOR_RGBA2RGB)
        else:
            # OpenCV读取的是BGR，转换为RGB
            img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)

        # 转换为灰度图
        gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)

        # 多种预处理方法
        # 1. 中值滤波去噪
        denoised = cv2.medianBlur(gray, 3)

        # 2. 使用自适应阈值二值化
        thresh_adaptive = cv2.adaptiveThreshold(denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)

        # 3. 使用Otsu阈值二值化
        _, thresh_otsu = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        # 4. 形态学操作（闭运算，填充内部孔洞；开运算，去除噪声）
        kernel = np.ones((2,2), np.uint8)
        closing = cv2.morphologyEx(thresh_otsu, cv2.MORPH_CLOSE, kernel)
        opening = cv2.morphologyEx(closing, cv2.MORPH_OPEN, kernel)

        # 返回所有预处理后的图像，供多个OCR引擎尝试
        return {
            'original': gray,
            'denoised': denoised,
            'thresh_adaptive': thresh_adaptive,
            'thresh_otsu': thresh_otsu,
            'morph': opening
        }

    def ocr_with_tesseract(self, image, lang='chi_sim'):
        """使用Tesseract进行OCR识别"""
        try:
            # 将numpy数组转换为PIL图像
            img_pil = Image.fromarray(image)
            # 配置Tesseract参数
            config = '--oem 3 --psm 6'
            text = pytesseract.image_to_string(img_pil, lang=lang, config=config)
            return text.strip()
        except Exception as e:
            print(f"Tesseract识别错误: {e}")
            return ""

    def ocr_with_easyocr(self, image):
        """使用EasyOCR进行OCR识别"""
        try:
            # EasyOCR需要numpy数组格式的图像
            results = self.easy_reader.readtext(image, paragraph=True)
            text = ' '.join([result[1] for result in results])
            return text.strip()
        except Exception as e:
            print(f"EasyOCR识别错误: {e}")
            return ""

    def recognize_text(self, image_path):
        """识别图像中的文字，结合多个OCR引擎的结果"""
        # 图像预处理
        processed_images = self.preprocess_image(image_path)

        best_text = ""
        best_confidence = 0

        # 尝试每种预处理图像
        for method, img in processed_images.items():
            print(f"尝试方法: {method}")

            # 使用Tesseract识别（中文）
            tesseract_text = self.ocr_with_tesseract(img, 'chi_sim')
            if tesseract_text:
                # 简单的置信度评估：文本长度（可以根据需要改进）
                confidence = len(tesseract_text)
                if confidence > best_confidence:
                    best_text = tesseract_text
                    best_confidence = confidence
                print(f"Tesseract 识别结果: {tesseract_text}")

            # 使用EasyOCR识别
            easyocr_text = self.ocr_with_easyocr(img)
            if easyocr_text:
                confidence = len(easyocr_text)
                if confidence > best_confidence:
                    best_text = easyocr_text
                    best_confidence = confidence
                print(f"EasyOCR 识别结果: {easyocr_text}")

        return best_text if best_text else "识别失败"

    def save_to_word(self, data, output_path):
        """将识别结果保存到Word文档"""
        doc = Document()
        doc.add_heading('OCR识别结果', 0)

        for filename, text in data:
            doc.add_heading(filename, level=1)
            doc.add_paragraph(text)
            doc.add_paragraph('-' * 50)

        doc.save(output_path)
        print(f"结果已保存到: {output_path}")

    def run(self):
        """运行OCR工具"""
        root = tk.Tk()
        root.withdraw()

        # 选择图片文件
        image_paths = filedialog.askopenfilenames(
            title="选择要识别的图片文件",
            filetypes=[("图片文件", "*.jpg *.jpeg *.png *.bmp *.tiff *.tif"), ("所有文件", "*.*")]
        )

        if not image_paths:
            return

        # 选择输出文件夹
        output_folder = filedialog.askdirectory(title="选择结果保存的文件夹")
        if not output_folder:
            return

        results = []
        for image_path in image_paths:
            filename = os.path.basename(image_path)
            print(f"处理中: {filename}")

            text = self.recognize_text(image_path)
            results.append((filename, text))

            print(f"识别结果: {text}\n")

        # 保存到Word文档
        output_path = os.path.join(output_folder, "OCR识别结果.docx")
        self.save_to_word(results, output_path)

        messagebox.showinfo("完成", f"处理完成！\n结果保存至: {output_path}")

if __name__ == "__main__":
    ocr_tool = ChineseOCR()
    ocr_tool.run()